#!/usr/bin/env python3
"""Generate all 18 eTables as a single Word document for BJA supplementary submission."""

import json, csv
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = Path(__file__).parent
RESULTS = BASE / 'results'


def load_json(name):
    with open(RESULTS / name) as f:
        return json.load(f)


def set_cell(cell, text, bold=False, size=8, align='left'):
    """Set cell text with formatting."""
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    run.bold = bold
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, title):
    """Add a formatted table to the document."""
    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(12)

    # Table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=8, align='center')
        # Gray background for header
        shading = table.rows[0].cells[i]._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'D9E2F3'
        })
        shading.append(shading_elem)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            align = 'left' if c_idx == 0 else 'center'
            set_cell(table.rows[r_idx + 1].cells[c_idx], val, size=8, align=align)

    doc.add_paragraph()  # spacing
    return table


def fmt_p(val, threshold=0.001):
    """Format p-value."""
    if val is None:
        return '—'
    if val < 0.001:
        return f'{val:.2e}'
    elif val < 0.05:
        return f'{val:.3f}'
    else:
        return f'{val:.3f}'


def fmt_ci(low, high, digits=2):
    return f'[{low:.{digits}f}, {high:.{digits}f}]'


# ======================================================================
# eTable generators
# ======================================================================

def etable1(doc):
    """eTable 1: NHD index RCS analysis."""
    d = load_json('rcs_RFTN_total_mcg_kg_NHD_pct.json')
    headers = ['Parameter', 'Value']
    rows = [
        ['n', f"{d['n']:,}"],
        ['Spline R²', f"{d['spline_r2']:.4f}"],
        ['Linear R²', f"{d['linear_r2']:.4f}"],
        ['P (nonlinear)', fmt_p(d['p_nonlinear'])],
        ['Knots', ', '.join(f"{k:.1f}" for k in d['knots'])],
        ['Interpretation', 'Non-significant nonlinear component (P = 0.079); monotonic increase in NHD with dose'],
    ]
    add_table(doc, headers, rows,
              'Table S1. Restricted Cubic Spline Analysis: Remifentanil Total Dose → NHD Index')


def etable2(doc):
    """eTable 2: IPTW on NHD index."""
    d = load_json('iptw_results.json')
    nhd = d['NHD_pct']
    meta = d['_meta']
    headers = ['Endpoint', 'Crude Δ', 'IPTW Δ', '95% CI', 'Significant']
    rows = [
        ['NHD index (%)', f"{nhd['crude_diff']:.3f}", f"{nhd['iptw_diff']:.3f}",
         fmt_ci(nhd['ci_low'], nhd['ci_high']), 'No'],
    ]
    add_table(doc, headers, rows,
              'Table S2. IPTW Analysis: NHD Index')
    note = (f"PS model fitted on n = 2,866 eligible cases (AUC = {meta['ps_auc']:.3f}); "
            f"PS trimming [0.05, 0.95] retained n = 2,584; "
            f"analytic n is endpoint-specific (n = {meta['n']:,} for NHD). "
            f"See Table S12 for full IPTW diagnostics.")
    p = doc.add_paragraph(f'Note: {note}')
    p.runs[0].font.size = Pt(7)
    p.runs[0].italic = True


def etable3(doc):
    """eTable 3: Partial vs complete OWSI sensitivity.

    Uses the authoritative per-endpoint RCS JSON files (same source as Figure 2
    and Table S14) so that n and P_nonlinear are consistent across the package.
    """
    d_partial  = load_json('rcs_RFTN_total_mcg_kg_OSI.json')        # partial OWSI
    d_complete = load_json('rcs_RFTN_total_mcg_kg_OWSI_complete.json')  # complete-case (Figure 2 Panel C)
    d_meta     = load_json('owsi_complete_case.json')  # only for correlation stat
    headers = ['Analysis', 'n', 'R²', 'R² (linear)', 'P (nonlinear)']
    rows = [
        ['OWSI partial (all)', f"{d_partial['n']:,}",
         f"{d_partial['spline_r2']:.4f}", f"{d_partial['linear_r2']:.4f}",
         fmt_p(d_partial['p_nonlinear'])],
        ['OWSI complete-case', f"{d_complete['n']:,}",
         f"{d_complete['spline_r2']:.4f}", f"{d_complete['linear_r2']:.4f}",
         fmt_p(d_complete['p_nonlinear'])],
    ]
    add_table(doc, headers, rows,
              'Table S3. OWSI Sensitivity: Partial vs Complete-Case Analysis')
    corr_r = d_meta['correlation_partial_complete']['r']
    p = doc.add_paragraph(
        f"Note: Correlation between partial and complete OWSI r = {corr_r:.4f}. "
        f"Complete-case analysis (n = {d_complete['n']:,}) corresponds to Figure 2 Panel C. "
        f"P_nonlinear = {fmt_p(d_complete['p_nonlinear'])} indicates the dose-response "
        f"curve does not exhibit significant departure from linearity in the complete-case sample."
    )
    p.runs[0].font.size = Pt(7)
    p.runs[0].italic = True


def etable4(doc):
    """eTable 4: Expanded covariates sensitivity."""
    d = load_json('expanded_covariates.json')
    headers = ['Model', 'n', 'R²', 'P (nonlinear)', 'IPTW Δ', '95% CI']
    rows = [
        ['Base covariates (RCS)', f"{d['rcs_original']['n']:,}",
         f"{d['rcs_original']['r2']:.4f}", fmt_p(d['rcs_original']['p_nonlinear']),
         f"{d['iptw_original']['iptw_diff']:.2f}",
         fmt_ci(d['iptw_original']['ci_low'], d['iptw_original']['ci_high'])],
        ['Expanded covariates (RCS)', f"{d['rcs_expanded']['n']:,}",
         f"{d['rcs_expanded']['r2']:.4f}", fmt_p(d['rcs_expanded']['p_nonlinear']),
         f"{d['iptw_expanded']['iptw_diff']:.2f}",
         fmt_ci(d['iptw_expanded']['ci_low'], d['iptw_expanded']['ci_high'])],
    ]
    comp = d['rcs_comparison']
    add_table(doc, headers, rows,
              'Table S4. Sensitivity Analysis: Expanded Covariate Adjustment')
    p = doc.add_paragraph(f"Note: R² Δ = {comp['r2_delta']:.4f}; IPTW change = {d['iptw_comparison']['change_pct']:.1f}%")
    p.runs[0].font.size = Pt(7)
    p.runs[0].italic = True


def etable5(doc):
    """eTable 5: Binary clinical endpoints."""
    d = load_json('binary_endpoints.json')
    headers = ['Endpoint', 'Events/n', 'Rate (%)', 'OR per 10 μg/kg', '95% CI', 'P']
    rows = []
    for key, label in [('HR_increase_20pct', 'HR > 20% increase'),
                       ('MAP_increase_20pct', 'MAP > 20% increase'),
                       ('Any_vasopressor', 'Any vasopressor use'),
                       ('Composite_event', 'Composite event')]:
        ep = d[key]
        lg = ep['logistic']['RFTN_total_mcg_kg']
        rows.append([
            label,
            f"{lg['n_events']}/{lg['n']}",
            f"{ep['overall_rate_pct']:.1f}",
            f"{lg['or']:.3f}",
            fmt_ci(lg['or_ci_low'], lg['or_ci_high']),
            fmt_p(lg['p_value']),
        ])
    add_table(doc, headers, rows,
              'Table S5. Binary Clinical Endpoints: Logistic Regression (Dose-Based)')

    # Taper-based
    headers2 = ['Endpoint', 'OR (taper slope)', '95% CI', 'P']
    rows2 = []
    for key, label in [('HR_increase_20pct', 'HR > 20% increase'),
                       ('MAP_increase_20pct', 'MAP > 20% increase'),
                       ('Any_vasopressor', 'Any vasopressor use'),
                       ('Composite_event', 'Composite event')]:
        lg = d[key]['logistic']['RFTN_taper_slope']
        rows2.append([label, f"{lg['or']:.3f}",
                      fmt_ci(lg['or_ci_low'], lg['or_ci_high']),
                      fmt_p(lg['p_value'])])
    add_table(doc, headers2, rows2,
              'Table S5 (continued). Binary Endpoints: Taper Slope Analysis')


def etable6(doc):
    """eTable 6: Taper dynamics expansion."""
    d = load_json('taper_expansion.json')
    tq = d['taper_quartile_analysis']
    headers = ['Outcome', 'Q1 (gradual)', 'Q2', 'Q3', 'Q4 (abrupt)', 'P (K-W)', 'ρ (trend)', 'P (trend)']
    rows = []
    for outcome in ['HR_rebound', 'MAP_rebound', 'OWSI']:
        t = tq[outcome]
        qm = t['quartile_means']
        label = outcome.replace('_', ' ')
        rows.append([
            label,
            f"{qm['Q1_gradual']:.2f} (n={t['quartile_ns']['Q1_gradual']})",
            f"{qm['Q2']:.2f} (n={t['quartile_ns']['Q2']})",
            f"{qm['Q3']:.2f} (n={t['quartile_ns']['Q3']})",
            f"{qm['Q4_abrupt']:.2f} (n={t['quartile_ns']['Q4_abrupt']})",
            fmt_p(t['kruskal_wallis_p']),
            f"{t['spearman_r_trend']:.3f}",
            fmt_p(t['p_trend']),
        ])
    add_table(doc, headers, rows,
              'Table S6. Taper Dynamics: Quartile Analysis by Taper Slope')


def etable7(doc):
    """eTable 7: Missing data report."""
    d = load_json('missingness_report.json')
    vm = d['variable_missingness']
    headers = ['Variable', 'n Available', 'n Missing', '% Missing']
    rows = []
    for var, info in sorted(vm.items(), key=lambda x: -x[1]['pct_missing']):
        if info['pct_missing'] > 0:
            rows.append([var, f"{info['n_available']:,}", f"{info['n_missing']:,}",
                         f"{info['pct_missing']:.1f}%"])
    add_table(doc, headers, rows,
              'Table S7. Missing Data Report')

    # Quartile missingness
    qm = d['quartile_missingness']
    headers2 = ['Variable', 'Q1 (%)', 'Q2 (%)', 'Q3 (%)', 'Q4 (%)', 'χ²', 'P']
    rows2 = []
    for var in ['HR_rebound', 'MAP_rebound', 'FTN_rescue_mcg_kg', 'NHD_pct',
                'RFTN_Ce_at_end', 'RFTN_taper_slope', 'intraop_ebl']:
        if var in qm:
            q = qm[var]
            pq = q['pct_by_quartile']
            chi = f"{q['chi2']:.1f}" if q['chi2'] is not None else '—'
            pv = fmt_p(q['p_value']) if q['p_value'] is not None else '—'
            rows2.append([var, f"{pq['Q1']:.1f}", f"{pq['Q2']:.1f}",
                          f"{pq['Q3']:.1f}", f"{pq['Q4']:.1f}", chi, pv])
    add_table(doc, headers2, rows2,
              'Table S7 (continued). Missingness by RFTN Dose Quartile')


def etable8(doc):
    """eTable 8: Volatile anesthetic sensitivity."""
    d = load_json('volatile_analysis.json')
    headers = ['Parameter', 'Value']
    hr = d['HR_rebound']
    rows = [
        ['TIVA n', f"{d['cohort']['TIVA']:,}"],
        ['Balanced n', f"{d['cohort']['balanced']:,}"],
        ['TIVA %', f"{d['cohort']['TIVA_pct']:.1f}%"],
        ['Base model β', f"{hr['base']['beta']:.5f}"],
        ['Base model P', fmt_p(hr['base']['p'])],
        ['Adjusted model β (+ anes_type)', f"{hr['adjusted']['beta']:.5f}"],
        ['Adjusted model P', fmt_p(hr['adjusted']['p'])],
        ['Anesthesia type β', f"{hr['adjusted']['anes_type_beta']:.3f}"],
        ['Anesthesia type P', fmt_p(hr['adjusted']['anes_type_p'])],
        ['Interaction P', f"{hr['interaction_p']:.4f}"],
        ['β change (%)', f"{hr['beta_change_pct']:.1f}%"],
    ]
    add_table(doc, headers, rows,
              'Table S8. Sensitivity Analysis: Volatile Anesthetic Adjustment')


def etable9(doc):
    """eTable 9: IPOW analysis."""
    d = load_json('reviewer2_analyses.json')
    ipow = d['major1_ipow']
    headers = ['Parameter', 'Value']
    diag = ipow['ipow_diagnostics']
    rows = [
        ['Observation model AUC', f"{ipow['observation_model']['auc']:.4f}"],
        ['Observation rate', f"{ipow['observation_model']['obs_rate']:.1%}"],
        ['n observed', f"{diag['n_observed']:,}"],
        ['n unobserved', f"{diag['n_unobserved']:,}"],
        ['Weight mean', f"{diag['weight_mean']:.3f}"],
        ['Weight median', f"{diag['weight_median']:.3f}"],
        ['Weight range', f"{diag['weight_min']:.3f} – {diag['weight_max']:.3f}"],
        ['ESS', f"{diag['ESS']:.1f}"],
    ]
    add_table(doc, headers, rows,
              'Table S9. Inverse Probability of Observation Weighting (IPOW) Diagnostics')

    # IPOW regression
    headers2 = ['Outcome', 'Complete-case β', 'P', 'IPOW β', 'P', 'Change (%)']
    rows2 = []
    hr = ipow['ipow_regression_HR']
    rows2.append(['HR rebound', f"{hr['complete_case_beta']:.5f}", fmt_p(hr['complete_case_p']),
                  f"{hr['ipow_beta']:.5f}", fmt_p(hr['ipow_p']),
                  f"{hr['pct_change']:.2f}%"])
    mp = ipow['ipow_regression_MAP']
    rows2.append(['MAP rebound', f"{mp['complete_case_beta']:.5f}", fmt_p(mp['complete_case_p']),
                  f"{mp['ipow_beta']:.5f}", fmt_p(mp['ipow_p']),
                  f"{mp.get('pct_change', 0):.2f}%"])
    add_table(doc, headers2, rows2,
              'Table S9 (continued). IPOW-Weighted Regression Results')


def etable10(doc):
    """eTable 10: Taper de-collinearity analysis."""
    d = load_json('reviewer2_analyses.json')
    t = d['major3_taper']
    # VIF
    headers = ['Variable', 'VIF']
    rows = [[k, f"{v:.3f}"] for k, v in t['vif'].items()]
    add_table(doc, headers, rows,
              'Table S10. Taper De-Collinearity: Variance Inflation Factors')

    # Partial correlations
    headers2 = ['Adjustment', 'ρ', 'P']
    for outcome in ['HR_rebound', 'MAP_rebound']:
        pc = t['partial_correlations'][outcome]
        rows2 = []
        for adj, label in [('unadjusted', 'Unadjusted'),
                           ('ctrl_rate', '+ Rate'),
                           ('ctrl_rate_Ce', '+ Rate + Ce'),
                           ('ctrl_rate_Ce_dose', '+ Rate + Ce + Dose')]:
            rows2.append([label, f"{pc[adj]['rho']:.4f}", fmt_p(pc[adj]['p'])])
        add_table(doc, headers2, rows2,
                  f'Table S10 (cont.). Partial Correlations: Taper → {outcome.replace("_", " ")}')

    # Elastic net
    en = t['elastic_net']
    headers3 = ['Variable', 'Standardized β', 'Rank']
    rows3 = []
    for rank, (var, coef) in enumerate(en['taper_rank'], 1):
        rows3.append([var, f"{coef:.4f}", str(rank)])
    add_table(doc, headers3, rows3,
              f'Table S10 (cont.). Elastic Net Feature Importance (n={en["n"]}, R²={en["r2"]:.4f})')


def etable11(doc):
    """eTable 11: Binary endpoints ARD/NNT."""
    d = load_json('reviewer2_analyses.json')
    ard = d['minor3_ard_nnt']
    headers = ['Endpoint', 'Q1 Risk', 'Q4 Risk', 'ARD', 'NNT']
    rows = []
    for key in ['HR_event_20', 'MAP_event_20']:
        a = ard[key]
        rows.append([a['label'], f"{a['Q1_risk']:.1%}", f"{a['Q4_risk']:.1%}",
                     f"{a['ARD']:.1%}", f"{a['NNT']:.0f}"])
    add_table(doc, headers, rows,
              'Table S11. Absolute Risk Differences and Numbers Needed to Treat')
    p = doc.add_paragraph(f"Note: ARD = absolute risk difference (Q4 − Q1); NNT = number needed to treat. n = {ard['HR_event_20']['n']:,}")
    p.runs[0].font.size = Pt(7)
    p.runs[0].italic = True


def etable12(doc):
    """eTable 12: Enhanced IPTW reporting."""
    d = load_json('reviewer2_analyses.json')
    iptw = d['minor4_iptw']['iptw_enhanced']
    headers = ['Parameter', 'Value']
    n_after_trim = iptw['n'] - iptw['ps_trimming']['n_trimmed']
    rows = [
        ['n total (PS-eligible)', f"{iptw['n']:,}"],
        ['n treated', f"{iptw['n_treated']:,}"],
        ['n control', f"{iptw['n_control']:,}"],
        ['PS model AUC', '0.851'],
        ['PS trimming threshold', f"{iptw['ps_trimming']['threshold']}"],
        ['n trimmed', f"{iptw['ps_trimming']['n_trimmed']}"],
        ['n after PS trimming (weight base)', f"{n_after_trim:,}"],
        ['Analytic n (hemodynamic endpoints)*', f'2,313 (subset of {n_after_trim:,} with post-emergence monitoring)'],
        ['Weight min', f"{iptw['weight_distribution']['min']:.3f}"],
        ['Weight median', f"{iptw['weight_distribution']['median']:.3f}"],
        ['Weight max', f"{iptw['weight_distribution']['max']:.1f}"],
        ['Weight mean (SD)', f"{iptw['weight_distribution']['mean']:.3f} ({iptw['weight_distribution']['sd']:.3f})"],
        ['ESS treated', f"{iptw['effective_sample_size']['ESS_treated']:.1f}"],
        ['ESS control', f"{iptw['effective_sample_size']['ESS_control']:.1f}"],
        ['ESS total', f"{iptw['effective_sample_size']['ESS_total']:.1f}"],
        ['ESS % of actual', f"{iptw['effective_sample_size']['ESS_pct_of_actual']:.1f}%"],
    ]
    add_table(doc, headers, rows,
              'Table S12. Enhanced IPTW Diagnostics\n'
              'Three-stage pipeline: (1) PS model fitted on n = 2,866 PS-eligible '
              'cases; (2) PS trimming [0.05, 0.95] removes 282 extreme propensity '
              f'cases \u2192 n = {n_after_trim:,} retained; (3) endpoint-specific outcome '
              'filtering \u2192 analytic n = 2,313 for hemodynamic endpoints (post-emergence '
              'monitoring required) or larger n for NHD/other endpoints with less missingness. '
              f'*Weight diagnostics computed on the full PS-trimmed sample (n = {n_after_trim:,}); '
              'PS AUC computed on PS-eligible sample (n = 2,866).')


def etable13(doc):
    """eTable 13: Comprehensive analysis grid."""
    d = load_json('reviewer2_analyses.json')
    grid = d['major4_grid']['analysis_grid']
    headers = ['ID', 'Description', 'n', 'Exposure', 'Endpoint', 'Covariates']
    rows = []
    for g in grid:
        rows.append([g['analysis_id'], g['description'], str(g['n']),
                     g['exposure'], g['endpoint'],
                     g['covariates'][:40] + ('...' if len(g['covariates']) > 40 else '')])
    add_table(doc, headers, rows,
              'Table S13. Comprehensive Analysis Grid')


def etable14(doc):
    """eTable 14: RCS statistical details for Figure 2.

    IMPORTANT: Must use the same RCS JSON files as Figure 2 Panels.
    Figure 2 Panel C = OWSI complete-case (n=2,847), NOT partial OSI (n=4,275).
    """
    rcs_files = [
        ('rcs_RFTN_total_mcg_kg_HR_rebound.json', 'HR rebound'),
        ('rcs_RFTN_total_mcg_kg_MAP_rebound.json', 'MAP rebound'),
        ('rcs_RFTN_total_mcg_kg_OWSI_complete.json', 'OWSI (complete-case)'),
    ]
    headers = ['Endpoint', 'n', 'Spline R²', 'Linear R²', 'P (nonlinear)', 'Knots']
    rows = []
    for fname, label in rcs_files:
        d = load_json(fname)
        rows.append([label, f"{d['n']:,}", f"{d['spline_r2']:.4f}",
                     f"{d['linear_r2']:.4f}", fmt_p(d['p_nonlinear']),
                     ', '.join(f"{k:.1f}" for k in d['knots'])])
    add_table(doc, headers, rows,
              'Table S14. RCS Statistical Details (Figure 2 Panels)')


def etable15(doc):
    """eTable 15: Rate threshold analysis."""
    d = load_json('rate_analysis_results.json')
    headers = ['Outcome', 'n', 'ρ (raw)', 'P (raw)', 'ρ (partial)', 'P (partial)', 'P (nonlinear)']
    rows = []
    for key in ['rate\u2192HR rebound', 'rate\u2192MAP rebound', 'rate\u2192FTN rescue',
                'rate\u2192OIH Surrogate Index', 'rate\u2192NHD index']:
        r = d[key]
        rows.append([key.replace('rate\u2192', ''),
                     f"{r['n']:,}",
                     f"{r['spearman_raw']:.3f}", fmt_p(r['p_raw']),
                     f"{r['partial_r']:.3f}", fmt_p(r['p_partial']),
                     fmt_p(r['rcs']['p_nonlinear'])])
    add_table(doc, headers, rows,
              'Table S15. Rate-Based Analysis: Correlations and RCS')

    # Quartile means
    headers2 = ['Outcome', 'Q1', 'Q2', 'Q3', 'Q4']
    rows2 = []
    for key in ['rate\u2192HR rebound', 'rate\u2192MAP rebound', 'rate\u2192FTN rescue',
                'rate\u2192OIH Surrogate Index', 'rate\u2192NHD index']:
        r = d[key]
        qm = r['quartile_means']
        rows2.append([key.replace('rate\u2192', '')] + [f"{v:.3f}" for v in qm])
    add_table(doc, headers2, rows2,
              'Table S15 (continued). Rate Quartile Means')


def etable16(doc):
    """eTable 16: Taper correlations."""
    d = load_json('taper_dynamics.json')
    headers = ['Correlation', 'n', 'ρ / Mean Diff', 'P']
    rows = []
    for key, info in d.items():
        if key.startswith('interaction'):
            rows.append([key, '—', f"β={info['beta']:.3f}", fmt_p(info['p'])])
        elif key.startswith('taper_type'):
            rows.append([key,
                         f"grad={info['n_gradual']}, abr={info['n_abrupt']}",
                         f"Δ={info['abrupt_mean'] - info['gradual_mean']:.2f}",
                         fmt_p(info['p'])])
        else:
            rows.append([key, f"{info['n']:,}", f"ρ={info['rho']:.3f}", fmt_p(info['p'])])
    add_table(doc, headers, rows,
              'Table S16. Taper Dynamics: Complete Correlation Matrix')


def etable17(doc):
    """eTable 17: Sensitivity analysis summary."""
    d = load_json('sensitivity_analyses.json')
    headers = ['Analysis', 'P (nonlinear)', 'Spline R²']
    rows = []
    for key, info in d.items():
        r2 = info.get('spline_r2', '—')
        r2_str = f"{r2:.4f}" if isinstance(r2, float) else r2
        rows.append([key, fmt_p(info['p_nonlinear']), r2_str])
    add_table(doc, headers, rows,
              'Table S17. Sensitivity Analysis Summary Across Exposure Metrics and Endpoints')


def etable18(doc):
    """eTable 18: Full baseline table by RFTN quartile."""
    headers = ['Variable', 'Q1', 'Q2', 'Q3', 'Q4', 'P']
    rows = []
    with open(RESULTS / 'table1_by_rftn_quartile.csv') as f:
        reader = csv.reader(f)
        next(reader)  # skip header
        for row in reader:
            if len(row) >= 5:
                rows.append(row[:6] if len(row) >= 6 else row[:5] + [''])
    add_table(doc, headers, rows,
              'Table S18. Full Baseline Characteristics by Remifentanil Dose Quartile')


# ======================================================================
# Main
# ======================================================================

def main():
    doc = Document()

    # Document title
    title = doc.add_heading('Supplementary Tables', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        'Remifentanil Discontinuation Dynamics and Immediate Emergence Hemodynamics:\n'
        'A Retrospective Analysis of 4,443 Noncardiac Surgeries from VitalDB'
    )
    run.font.size = Pt(10)
    run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing

    # Generate all 18 eTables
    generators = [
        etable1, etable2, etable3, etable4, etable5, etable6,
        etable7, etable8, etable9, etable10, etable11, etable12,
        etable13, etable14, etable15, etable16, etable17, etable18,
    ]

    for i, gen in enumerate(generators, 1):
        print(f"  Generating eTable {i}...")
        gen(doc)
        if i < 18:
            doc.add_page_break()

    # Save
    outpath = BASE / 'supplementary_tables' / 'Supplementary_Tables_S1-S18.docx'
    doc.save(str(outpath))
    print(f"\nSaved: {outpath}")
    print(f"Total eTables: 18")


if __name__ == '__main__':
    main()
