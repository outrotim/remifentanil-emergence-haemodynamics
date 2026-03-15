#!/usr/bin/env python3
"""Generate main text Tables 1 & 2 for BJA submission."""

import json, csv
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = Path(__file__).parent
RESULTS = BASE / 'results'


def load_json(name):
    with open(RESULTS / name) as f:
        return json.load(f)


def set_cell(cell, text, bold=False, size=9, align='left'):
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    run.bold = bold
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p.paragraph_format.space_before = Pt(12)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=9, align='center')
        shading = table.rows[0].cells[i]._element.get_or_add_tcPr()
        elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'D9E2F3'
        })
        shading.append(elem)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            align = 'left' if c_idx == 0 else 'center'
            set_cell(table.rows[r_idx + 1].cells[c_idx], val, size=9, align=align)

    doc.add_paragraph()
    return table


def fmt_p(val):
    if val is None:
        return '\u2014'
    if val < 0.001:
        return '<0.001'
    return f'{val:.3f}'


def fmt_ci(low, high, d=2):
    return f'[{low:.{d}f}, {high:.{d}f}]'


def iptw_p_from_ci(estimate, ci_low, ci_high):
    """Approximate P value from IPTW point estimate and bootstrap CI."""
    from scipy import stats
    se = (ci_high - ci_low) / 3.92  # 95% CI width / 2*1.96
    if se <= 0:
        return 1.0
    z = abs(estimate / se)
    return float(2 * stats.norm.sf(z))


def table1(doc):
    """Table 1: Baseline characteristics by RFTN quartile."""
    headers = ['Variable', 'Q1 (n \u2248 1,093)', 'Q2 (n \u2248 1,093)', 'Q3 (n \u2248 1,093)', 'Q4 (n \u2248 1,093)', 'P']
    rows = []
    with open(RESULTS / 'table1_by_rftn_quartile.csv') as f:
        reader = csv.reader(f)
        next(reader)
        for row in reader:
            if len(row) >= 5:
                # Skip [Outcome] rows \u2014 outcomes belong in Table 2
                if row[0].startswith('[Outcome]'):
                    continue
                p_val = row[5] if len(row) >= 6 else ''
                rows.append([row[0], row[1], row[2], row[3], row[4], p_val])
    add_table(doc, headers, rows,
              'Table 1. Baseline and Operative Characteristics by Remifentanil Dose Quartile')
    note = doc.add_paragraph()
    run = note.add_run(
        'Values are mean (SD) for continuous variables and n (%) for categorical variables. '
        'P values from Kruskal-Wallis test (continuous) or \u03c7\u00b2 test (categorical). '
        'Quartile ranges: Q1 \u2264 9.8, Q2 9.8\u201315.8, Q3 15.8\u201326.4, Q4 > 26.4 \u03bcg/kg.'
    )
    run.font.size = Pt(8)
    run.italic = True


def table2(doc):
    """Table 2: Composite key effects."""
    doc.add_page_break()

    # Panel A: Quartile effects (dose quartiles \u2014 consistent with Table 1 and Section 3.4)
    dose_q = load_json('dose_quartile_analysis_results.json')
    headers_a = ['Outcome', 'Q1', 'Q2', 'Q3', 'Q4', 'P (trend)']
    rows_a = []
    for key, label in [('dose\u2192HR rebound (bpm)', 'HR rebound (bpm)'),
                       ('dose\u2192MAP rebound (mmHg)', 'MAP rebound (mmHg)'),
                       ('dose\u2192FTN rescue (\u03bcg/kg)', 'FTN rescue (\u03bcg/kg)'),
                       ('dose\u2192OWSI (Z-score)', 'OWSI (Z-score)')]:
        r = dose_q[key]
        qm = r['quartile_means']
        rows_a.append([label] + [f"{v:.2f}" for v in qm] + [fmt_p(float(r['p_trend']) if r['p_trend'] != 0 else 0)])
    add_table(doc, headers_a, rows_a,
              'Table 2. Composite Key Effects')
    p = doc.add_paragraph(
        'Panel A. Outcome means by remifentanil total dose quartile '
        '(Q1 \u2264 9.8, Q2 9.8\u201315.8, Q3 15.8\u201326.4, Q4 > 26.4 \u03bcg/kg; '
        'consistent with Table 1 stratification).')
    p.runs[0].font.size = Pt(8)
    p.runs[0].italic = True

    # Panel B: IPTW ATE
    iptw = load_json('iptw_results.json')
    headers_b = ['Outcome', 'Crude \u0394', 'IPTW ATE', '95% CI', 'P']
    rows_b = []
    for key, label in [('HR_rebound', 'HR rebound (bpm)'),
                       ('MAP_rebound', 'MAP rebound (mmHg)'),
                       ('FTN_rescue_mcg_kg', 'FTN rescue (\u03bcg/kg)'),
                       ('OSI', 'OWSI (Z-score)'),
                       ('NHD_pct', 'NHD index (%)')]:
        r = iptw[key]
        p_iptw = iptw_p_from_ci(r['iptw_diff'], r['ci_low'], r['ci_high'])
        rows_b.append([label, f"{r['crude_diff']:.2f}", f"{r['iptw_diff']:.2f}",
                       fmt_ci(r['ci_low'], r['ci_high']),
                       fmt_p(p_iptw)])
    add_table(doc, headers_b, rows_b, '')
    p = doc.add_paragraph(
        f"Panel B. IPTW Average Treatment Effect (High vs Low infusion rate). "
        f"PS model fitted on n = 2,866 eligible cases (PS AUC = {iptw['_meta']['ps_auc']:.3f}); "
        f"PS trimming [0.05, 0.95] retained n = 2,584; "
        f"analytic n is endpoint-specific (e.g., n = {iptw['_meta']['n']:,} for hemodynamic outcomes "
        f"with post-emergence monitoring; NHD uses a larger n). "
        f"Treatment = above-median mean infusion rate "
        f"(threshold = {iptw['_meta']['rate_threshold']:.3f} \u03bcg/kg/min); "
        f"mean post-weighting SMD = {iptw['_meta']['mean_smd_after']:.3f}. "
        f"P values from bootstrap 95% CI via Wald method. "
        f"See Table S12 for full diagnostics.")
    p.runs[0].font.size = Pt(8)
    p.runs[0].italic = True

    # Panel C: Taper dynamics
    taper = load_json('taper_dynamics.json')
    headers_c = ['Comparison', 'n', '\u03c1 / \u0394', 'P']
    rows_c = []
    for key, label in [('RFTN_taper_slope\u2192HR_rebound', 'Taper \u2192 HR rebound'),
                        ('RFTN_taper_slope\u2192MAP_rebound', 'Taper \u2192 MAP rebound'),
                        ('RFTN_taper_slope\u2192OWSI_complete', 'Taper \u2192 OWSI (cc)')]:
        t = taper[key]
        rows_c.append([label, f"{t['n']:,}", f"\u03c1 = {t['rho']:.3f}", fmt_p(t['p'])])

    tt_hr = taper['taper_type\u2192HR_rebound']
    rows_c.append(['Above vs Below median \u2192 HR',
                   f"{tt_hr['n_gradual'] + tt_hr['n_abrupt']:,}",
                   f"\u0394 = {tt_hr['abrupt_mean'] - tt_hr['gradual_mean']:.2f} bpm",
                   fmt_p(tt_hr['p'])])
    tt_osi = taper['taper_type\u2192OWSI_complete']
    rows_c.append(['Above vs Below median \u2192 OWSI (cc)',
                   f"{tt_osi['n_gradual'] + tt_osi['n_abrupt']:,}",
                   f"\u0394 = {tt_osi['abrupt_mean'] - tt_osi['gradual_mean']:.2f}",
                   fmt_p(tt_osi['p'])])
    add_table(doc, headers_c, rows_c, '')
    p = doc.add_paragraph(
        'Panel C. Taper dynamics: Spearman correlations (rows 1\u20133) and median-split '
        'comparison (rows 4\u20135) of taper slope vs hemodynamic rebound. '
        'Median split at taper slope median within the analysis subset; binary groups restricted to '
        'cases with non-missing taper slope. '
        'OWSI (cc) = Opioid Withdrawal Surrogate Index, complete-case (requires all three '
        'components: HR rebound, MAP rebound, fentanyl rescue). '
        'Spearman n reflects endpoint-specific pairwise complete observations; '
        'binary n reflects taper-slope-available \u00d7 endpoint-available intersection. '
        'See Table S16 for partial OWSI results (n \u2248 4,058).')
    p.runs[0].font.size = Pt(8)
    p.runs[0].italic = True


def main():
    doc = Document()

    title = doc.add_heading('Main Text Tables', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table1(doc)
    table2(doc)

    outpath = BASE / 'main_figures' / 'Main_Tables_1-2.docx'
    doc.save(str(outpath))
    print(f"Saved: {outpath}")


if __name__ == '__main__':
    main()
